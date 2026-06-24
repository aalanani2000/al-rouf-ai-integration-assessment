from __future__ import annotations

import os
import subprocess
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    Image as PdfImage,
    ListFlowable,
    ListItem,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)


ROOT = Path(__file__).resolve().parents[1]
REPORT_1_DIR = ROOT / "reports" / "01_execution_evidence_report"
REPORT_2_DIR = ROOT / "reports" / "02_final_result_report"
SCREENSHOT_DIR = ROOT / "docs" / "screenshots"
RENDERER = Path(
    r"C:\Users\idhoo\.codex\plugins\cache\openai-primary-runtime\documents\26.513.11550"
    r"\skills\documents\render_docx.py"
)
PYTHON = Path(
    r"C:\Users\idhoo\.cache\codex-runtimes\codex-primary-runtime\dependencies\python\python.exe"
)

REPO_URL = "https://github.com/aalanani2000/al-rouf-ai-integration-assessment"
DEFAULT_BRANCH = "main"
LATEST_COMMIT = "653ec552871083dca14498d120d6e6c156ed4f31"
REPORT_DATE = date(2026, 6, 24).isoformat()

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(34, 34, 34)
MUTED = RGBColor(90, 90, 90)
HEADER_FILL = "F2F4F7"


def repo_tree() -> str:
    return """C:.
|   .env
|   .env.example
|   .gitignore
|   docker-compose.yml
|   README.md
|
+---data
|       alerts.json
|       mock_crm.json
|
+---docs
|   +---architecture
|   +---evidence
|   |   \\---task1_rfq_crm_automation
|   |           execution_notes.md
|   |
|   \\---screenshots
|       +---Task 1
|       |       javascript code 1 error .png
|       |       n8n canva.png
|       |       open ai node.png
|       |       proof of the volumes being correctly attached.png
|       |       repository structure.png
|       |       set field output.png
|       |       set field.png
|       |       Webhook node.png
|       |       workflow progress .png
|       |       workflow progress 1 .png
|       |       workflow succes.png
|       |       workflow.png
|       |       workflow2.png
|       |
|       +---Task 2
|       |       Open API.png
|       |
|       \\---Task 3
|               a problem when communicating with model.png
|               testing the RAG on arabic question.pdf
|
+---reports
|   +---01_execution_evidence_report
|   \\---02_final_result_report
+---task1_rfq_crm_automation
|   \\---n8n_workflows
|           rfq-crm-automation.json
|
+---task2_quotation_microservice
|   |   .gitignore
|   |   docker-compose.yml
|   |   Dockerfile
|   |   README.md
|   |   requirements.txt
|   |
|   +---app
|   |       catalog.py
|   |       database.py
|   |       main.py
|   |       quoting.py
|   |       schemas.py
|   |       __init__.py
|   |
|   +---data
|   |       quotes.db
|   |
|   \\---tests
|           test_quotation.py
|           __init__.py
|
\\---task3_bilingual_rag
    |   .env
    |   .env.example
    |   .gitignore
    |   docker-compose.yml
    |   Dockerfile
    |   README.md
    |   requirements.txt
    |
    +---app
    |       embeddings.py
    |       llm.py
    |       loader.py
    |       main.py
    |       schemas.py
    |       __init__.py
    |
    +---data
    |       embeddings_cache.json
    |
    +---documents
    |       doc1_product_spec_high_bay_en.md
    |       doc2_warranty_policy_ar.md
    |       doc3_shipping_terms_en.md
    |       doc4_product_spec_street_light_ar.md
    |
    \\---tests
            test_rag.py
            __init__.py"""


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False, color: RGBColor | None = None) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(10)
    if color:
        run.font.color.rgb = color
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def style_table(table) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for cell in hdr.cells:
        set_cell_shading(cell, HEADER_FILL)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.color.rgb = DARK_BLUE
    for row in table.rows:
        for cell in row.cells:
            cell.margin_top = 80
            cell.margin_bottom = 80
            cell.margin_left = 120
            cell.margin_right = 120


def setup_doc(title: str, subtitle: str) -> Document:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(11)
    styles["Normal"].font.color.rgb = INK
    styles["Normal"].paragraph_format.space_after = Pt(6)
    styles["Normal"].paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    header_p = section.header.paragraphs[0]
    header_p.text = "AL ROUF LED AI Integration Assessment"
    header_p.runs[0].font.size = Pt(9)
    header_p.runs[0].font.color.rgb = MUTED

    footer_p = section.footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_p.text = f"Prepared {REPORT_DATE}"
    footer_p.runs[0].font.size = Pt(9)
    footer_p.runs[0].font.color.rgb = MUTED

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(23)
    run.font.color.rgb = RGBColor(0, 0, 0)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(16)
    run = p.add_run(subtitle)
    run.font.size = Pt(13)
    run.font.color.rgb = MUTED

    meta = doc.add_table(rows=5, cols=2)
    meta.columns[0].width = Inches(1.3)
    meta.columns[1].width = Inches(5.2)
    rows = [
        ("Repository", REPO_URL),
        ("Default branch", DEFAULT_BRANCH),
        ("Latest commit", LATEST_COMMIT),
        ("Prepared date", REPORT_DATE),
        ("Assessment", "AI Integration Engineer task pack"),
    ]
    for row, (label, value) in zip(meta.rows, rows):
        set_cell_text(row.cells[0], label, bold=True, color=DARK_BLUE)
        set_cell_text(row.cells[1], value)
    style_table(meta)
    return doc


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def add_status_table(doc: Document) -> None:
    table = doc.add_table(rows=1, cols=4)
    headers = ["Task", "Implementation", "Validation", "Status"]
    for cell, header in zip(table.rows[0].cells, headers):
        set_cell_text(cell, header, bold=True, color=DARK_BLUE)
    rows = [
        (
            "Task 1",
            "n8n RFQ webhook, OpenAI extraction, mock CRM, bilingual reply, alert log.",
            "Successful n8n executions, persisted JSON records, screenshot evidence.",
            "Complete; attachment bytes scoped as future enhancement.",
        ),
        (
            "Task 2",
            "FastAPI quotation API with catalog, pricing tiers, SQLite, Docker, OpenAPI.",
            "Automated pytest suite and OpenAPI screenshot evidence.",
            "Complete.",
        ),
        (
            "Task 3",
            "FastAPI bilingual RAG service over English/Arabic documents with citations and refusal.",
            "Unit tests, integration test path, Arabic question evidence, error-handling screenshot.",
            "Complete.",
        ),
    ]
    for values in rows:
        cells = table.add_row().cells
        for cell, text in zip(cells, values):
            set_cell_text(cell, text)
    style_table(table)


def image_width(path: Path, max_width: float = 6.2, max_height: float = 4.2) -> float:
    with Image.open(path) as img:
        width, height = img.size
    ratio = width / max(height, 1)
    if ratio >= max_width / max_height:
        return max_width
    return min(max_width, max_height * ratio)


def add_image(doc: Document, path: Path, caption: str) -> None:
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    p.add_run().add_picture(str(path), width=Inches(image_width(path)))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(10)
    run = cap.add_run(caption)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED


def build_evidence_report() -> Path:
    doc = setup_doc(
        "01 Execution Evidence Report",
        "Step-by-step build evidence, screenshots, validation, decisions, and ownership disclosure.",
    )

    doc.add_heading("Assessment Requirements Covered", level=1)
    add_bullets(
        doc,
        [
            "Runnable GitHub repository with setup documentation, task folders, screenshots, and report deliverables.",
            "RFQ to CRM automation using n8n as an equivalent orchestration approach.",
            "FastAPI quotation microservice with tests, Docker packaging, and OpenAPI documentation.",
            "Bilingual RAG workflow over sample English and Arabic documents with citations and refusal behavior.",
            "Explicit notes on error handling, maintainability, security hygiene, trade-offs, validation, and AI assistance.",
        ],
    )

    doc.add_heading("Work Evidence By Task", level=1)
    add_status_table(doc)

    doc.add_heading("Task 1 Evidence: RFQ To CRM Automation", level=1)
    add_numbered(
        doc,
        [
            "Created a Docker-based n8n environment with a named n8n data volume and a bind-mounted local data folder.",
            "Built a webhook intake endpoint for RFQ-style JSON payloads.",
            "Configured an OpenAI node to extract urgency, key requirements, and a concise summary from unstructured RFQ text.",
            "Merged structured webhook fields with extracted intent and generated CRM metadata.",
            "Persisted records into a mock CRM JSON file and wrote sales alerts with bilingual reply drafts.",
            "Debugged n8n sandbox limitations, missing-file behavior, and node-reference mismatches during testing.",
        ],
    )
    add_bullets(
        doc,
        [
            "Key trade-off: webhook simulation was used instead of email/IMAP to keep the workflow reproducible without OAuth setup.",
            "Key limitation: attachment filenames are captured as metadata; real file-byte archiving is documented as a production enhancement.",
            "Validation: multiple full workflow executions produced successful CRM and alert outputs.",
        ],
    )
    task1_images = [
        "repository structure.png",
        "proof of the volumes being correctly attached.png",
        "Webhook node.png",
        "n8n canva.png",
        "open ai node.png",
        "set field.png",
        "set field output.png",
        "javascript code 1 error .png",
        "workflow progress .png",
        "workflow progress 1 .png",
        "workflow.png",
        "workflow2.png",
        "workflow succes.png",
    ]
    for name in task1_images:
        add_image(doc, SCREENSHOT_DIR / "Task 1" / name, f"Task 1 evidence: {name}")

    doc.add_heading("Task 2 Evidence: Quotation Microservice", level=1)
    add_bullets(
        doc,
        [
            "Built a FastAPI service with `/health`, `/catalog`, `POST /quote`, and `GET /quote/{quote_id}` endpoints.",
            "Implemented product catalog and bulk-discount pricing logic with offline-friendly mocked data.",
            "Persisted quotes in SQLite and packaged the service with Docker.",
            "Added automated tests for health, catalog, quote creation, discount tiers, validation failures, and quote retrieval.",
        ],
    )
    add_image(doc, SCREENSHOT_DIR / "Task 2" / "Open API.png", "Task 2 evidence: OpenAPI documentation page.")

    doc.add_heading("Task 3 Evidence: Bilingual RAG Knowledge Workflow", level=1)
    add_bullets(
        doc,
        [
            "Built a FastAPI service that loads 4 sample documents in English and Arabic.",
            "Embedded document chunks with OpenAI embeddings and cached them locally.",
            "Retrieved relevant chunks using cosine similarity and refused out-of-scope questions before generation.",
            "Returned same-language answers with citations, latency, and estimated cost fields.",
            "Tested Arabic handling, retrieval behavior, refusal logic, and optional OpenAI-backed integration paths.",
        ],
    )
    add_image(
        doc,
        SCREENSHOT_DIR / "Task 3" / "a problem when communicating with model.png",
        "Task 3 evidence: model communication issue captured during testing.",
    )
    doc.add_paragraph(
        "Additional Task 3 evidence is stored at docs/screenshots/Task 3/testing the RAG on arabic question.pdf.",
    )

    doc.add_heading("Validation Summary", level=1)
    add_bullets(
        doc,
        [
            "Task 1 was validated through end-to-end n8n executions and persisted mock output files.",
            "Task 2 includes a pytest suite documented in its README; OpenAPI evidence confirms the interactive API surface.",
            "Task 3 includes offline tests for document loading, chunking, search, and refusal behavior, plus integration tests gated by `OPENAI_API_KEY`.",
            "Docker packaging is included for all runtime services that require containers.",
        ],
    )

    doc.add_heading("Ownership And AI Assistance Disclosure", level=1)
    doc.add_paragraph(
        "AI assistance was used as a technical mentor and documentation partner for architecture discussion, starter logic, debugging support, and final writing polish. The candidate personally completed the hands-on implementation work: n8n node setup, Docker execution, API implementation, test execution, screenshot collection, environment configuration, validation, and final engineering decisions."
    )

    REPORT_1_DIR.mkdir(parents=True, exist_ok=True)
    path = REPORT_1_DIR / "01_execution_evidence_report.docx"
    doc.save(path)
    return path


def build_final_report() -> Path:
    doc = setup_doc(
        "02 Final Result Report",
        "Final outputs, architecture summary, setup guide, API evidence, limitations, and improvements.",
    )

    doc.add_heading("Final Outcome", level=1)
    doc.add_paragraph(
        "The assessment implementation is complete across all three requested tasks. The repository contains runnable source code, Docker packaging, tests where required, exported workflow artifacts, setup documentation, screenshots, and the final reports."
    )
    add_status_table(doc)

    doc.add_heading("Task 1 Final Result", level=1)
    add_bullets(
        doc,
        [
            "Inbound RFQ data enters through an n8n webhook.",
            "OpenAI extracts urgency, key requirements, and a summary from the free-text message.",
            "The workflow creates a mock CRM record in `data/mock_crm.json`.",
            "The workflow drafts English and Arabic reply text without including pricing.",
            "The workflow creates an internal sales alert in `data/alerts.json`.",
        ],
    )

    doc.add_heading("Task 2 Final Result", level=1)
    add_bullets(
        doc,
        [
            "FastAPI service exposes health, catalog, quote creation, and quote retrieval endpoints.",
            "Bulk discount tiers are applied by total quote quantity.",
            "SQLite stores generated quotes and survives restarts when Docker volume/data mapping is used.",
            "OpenAPI documentation is available at `/docs` while the service is running.",
            "Automated tests cover expected success and failure paths.",
        ],
    )

    doc.add_heading("Task 3 Final Result", level=1)
    add_bullets(
        doc,
        [
            "The service indexes 4 sample documents covering product specifications, warranty policy, and shipping terms.",
            "English and Arabic questions can retrieve relevant source chunks.",
            "Responses cite source documents and include latency and estimated cost metadata.",
            "Out-of-scope questions are refused when retrieval confidence is insufficient.",
            "Embeddings are cached locally to avoid unnecessary repeated embedding calls.",
        ],
    )

    doc.add_heading("Architecture Explanation", level=1)
    add_bullets(
        doc,
        [
            "Task 1 uses n8n for orchestration because the workflow itself can be exported, versioned, reviewed, and rerun locally.",
            "Task 2 separates API schemas, catalog data, pricing logic, database access, and route handlers for maintainability.",
            "Task 3 keeps the RAG system intentionally lightweight because the assessment document set is small; a full vector database would be unnecessary infrastructure at this scale.",
        ],
    )

    doc.add_heading("Run And Test Commands", level=1)
    commands = doc.add_table(rows=1, cols=3)
    for cell, header in zip(commands.rows[0].cells, ["Area", "Command", "Purpose"]):
        set_cell_text(cell, header, bold=True, color=DARK_BLUE)
    rows = [
        ("Task 1", "docker-compose up -d", "Start n8n from the repository root."),
        ("Task 2", "cd task2_quotation_microservice && docker-compose up --build", "Run quotation API on port 8000."),
        ("Task 2 tests", "cd task2_quotation_microservice && PYTHONPATH=. pytest tests/ -v", "Run quotation service tests."),
        ("Task 3", "cd task3_bilingual_rag && docker-compose up --build", "Run bilingual RAG API on port 8001."),
        ("Task 3 tests", "cd task3_bilingual_rag && PYTHONPATH=. pytest tests/ -v", "Run RAG tests; API-key integration tests skip when unset."),
    ]
    for row in rows:
        cells = commands.add_row().cells
        for cell, text in zip(cells, row):
            set_cell_text(cell, text)
    style_table(commands)

    doc.add_heading("Repository Guidance", level=1)
    doc.add_paragraph("Use this structure map to navigate the solution bundle and GitHub repository:")
    tree_para = doc.add_paragraph()
    tree_para.style = doc.styles["Normal"]
    tree_para.paragraph_format.line_spacing = 1.0
    tree_para.paragraph_format.space_after = Pt(8)
    run = tree_para.add_run(repo_tree())
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Consolas")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Consolas")
    run.font.size = Pt(7.5)

    doc.add_heading("Error Handling, Maintainability, And Security", level=1)
    add_bullets(
        doc,
        [
            "Task 1 persists n8n state and mock CRM data with Docker volumes and local bind mounts.",
            "Task 2 validates input with Pydantic, returns clear HTTP errors, and avoids leaking internal stack traces to clients.",
            "Task 3 refuses unsupported questions through retrieval thresholding rather than relying only on prompt instructions.",
            "Secrets are stored in `.env` files or runtime environment variables and are not hardcoded in source files.",
            "The FastAPI codebases keep core logic separate from HTTP handlers so behavior can be tested directly.",
        ],
    )

    doc.add_heading("Future Enhancements", level=1)
    add_bullets(
        doc,
        [
            "Task 1: ingest RFQs from real email inboxes, WhatsApp Business, website forms, or CRM webhooks instead of only a simulated webhook.",
            "Task 1: archive actual attachment files under each CRM record and add duplicate detection.",
            "Task 2: add a quotation UI, quote PDF generation, authentication, configurable product/pricing administration, taxes, shipping, and ERP/CRM integration.",
            "Task 3: add a polished bilingual chat UI, document upload and re-indexing, richer citation snippets, streaming responses, access control, and a vector database if the corpus grows.",
        ],
    )

    doc.add_heading("AI Assistance Disclosure", level=1)
    doc.add_paragraph(
        "AI assistance was used for brainstorming architecture choices, debugging explanations, starter snippets, and documentation/report polish. The candidate remained responsible for all local execution, implementation, testing, configuration, screenshots, trade-off selection, and final submission packaging."
    )

    REPORT_2_DIR.mkdir(parents=True, exist_ok=True)
    path = REPORT_2_DIR / "02_final_result_report.docx"
    doc.save(path)
    return path


def render_docx(path: Path) -> None:
    out_dir = path.parent / "_rendered"
    out_dir.mkdir(parents=True, exist_ok=True)
    try:
        subprocess.run(
            [str(PYTHON), str(RENDERER), str(path), "--output_dir", str(out_dir), "--emit_pdf"],
            cwd=str(ROOT),
            check=True,
            capture_output=True,
            text=True,
        )
        pdf = out_dir / f"{path.stem}.pdf"
        if pdf.exists():
            target = path.with_suffix(".pdf")
            target.write_bytes(pdf.read_bytes())
            return
    except Exception:
        print(f"DOCX render skipped for {path.name}: LibreOffice/soffice is not available; creating direct PDF fallback.")
        try:
            out_dir.rmdir()
        except OSError:
            pass

    if path.stem == "01_execution_evidence_report":
        build_evidence_pdf(path.with_suffix(".pdf"))
    elif path.stem == "02_final_result_report":
        build_final_pdf(path.with_suffix(".pdf"))


def pdf_styles():
    styles = getSampleStyleSheet()
    styles.add(
        ParagraphStyle(
            "TitleBlack",
            parent=styles["Title"],
            fontName="Helvetica-Bold",
            fontSize=22,
            leading=26,
            textColor=colors.black,
            spaceAfter=8,
        )
    )
    styles.add(
        ParagraphStyle(
            "SubtitleMuted",
            parent=styles["Normal"],
            fontName="Helvetica",
            fontSize=11,
            leading=15,
            textColor=colors.HexColor("#5A5A5A"),
            spaceAfter=14,
        )
    )
    styles.add(
        ParagraphStyle(
            "H1Blue",
            parent=styles["Heading1"],
            fontName="Helvetica-Bold",
            fontSize=15,
            leading=18,
            textColor=colors.HexColor("#2E74B5"),
            spaceBefore=12,
            spaceAfter=7,
        )
    )
    styles.add(
        ParagraphStyle(
            "Caption",
            parent=styles["Normal"],
            fontName="Helvetica-Oblique",
            fontSize=8.5,
            leading=11,
            textColor=colors.HexColor("#5A5A5A"),
            alignment=TA_CENTER,
            spaceAfter=8,
        )
    )
    styles.add(
        ParagraphStyle(
            "CodeSmall",
            parent=styles["Code"],
            fontName="Courier",
            fontSize=6.6,
            leading=7.5,
            spaceAfter=8,
        )
    )
    return styles


def pdf_doc(path: Path):
    return SimpleDocTemplate(
        str(path),
        pagesize=letter,
        rightMargin=0.72 * inch,
        leftMargin=0.72 * inch,
        topMargin=0.72 * inch,
        bottomMargin=0.72 * inch,
        title=path.stem,
    )


def pdf_meta_table(styles):
    data = [
        ["Repository", REPO_URL],
        ["Default branch", DEFAULT_BRANCH],
        ["Latest commit", LATEST_COMMIT],
        ["Prepared date", REPORT_DATE],
        ["Assessment", "AI Integration Engineer task pack"],
    ]
    table = Table(data, colWidths=[1.25 * inch, 5.15 * inch])
    table.setStyle(
        TableStyle(
            [
                ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#D0D7DE")),
                ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#F2F4F7")),
                ("TEXTCOLOR", (0, 0), (0, -1), colors.HexColor("#1F4D78")),
                ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"),
                ("FONTNAME", (1, 0), (1, -1), "Helvetica"),
                ("FONTSIZE", (0, 0), (-1, -1), 8.5),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("LEFTPADDING", (0, 0), (-1, -1), 6),
                ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ]
        )
    )
    return table


def pdf_bullets(styles, items: list[str]):
    return ListFlowable(
        [ListItem(Paragraph(item, styles["BodyText"]), leftIndent=12) for item in items],
        bulletType="bullet",
        leftIndent=18,
    )


def pdf_numbered(styles, items: list[str]):
    return ListFlowable(
        [ListItem(Paragraph(item, styles["BodyText"]), leftIndent=12) for item in items],
        bulletType="1",
        leftIndent=18,
    )


def pdf_status_table(styles):
    data = [
        ["Task", "Implementation", "Validation", "Status"],
        [
            "Task 1",
            "n8n RFQ webhook, OpenAI extraction, mock CRM, bilingual reply, alert log.",
            "Successful n8n executions, persisted JSON records, screenshot evidence.",
            "Complete; attachment bytes scoped as future enhancement.",
        ],
        [
            "Task 2",
            "FastAPI quotation API with catalog, pricing tiers, SQLite, Docker, OpenAPI.",
            "Automated pytest suite and OpenAPI screenshot evidence.",
            "Complete.",
        ],
        [
            "Task 3",
            "FastAPI bilingual RAG service over English/Arabic documents with citations and refusal.",
            "Unit tests, Arabic question evidence, error-handling screenshot.",
            "Complete.",
        ],
    ]
    wrapped = [[Paragraph(str(cell), styles["BodyText"]) for cell in row] for row in data]
    table = Table(wrapped, colWidths=[0.75 * inch, 2.0 * inch, 2.2 * inch, 1.45 * inch], repeatRows=1)
    table.setStyle(
        TableStyle(
            [
                ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#D0D7DE")),
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#F2F4F7")),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#1F4D78")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("FONTSIZE", (0, 0), (-1, -1), 8),
                ("LEFTPADDING", (0, 0), (-1, -1), 5),
                ("RIGHTPADDING", (0, 0), (-1, -1), 5),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ]
        )
    )
    return table


def pdf_image_flow(path: Path, caption: str, styles):
    if not path.exists():
        return []
    with Image.open(path) as img:
        width, height = img.size
    max_w = 6.4 * inch
    max_h = 4.2 * inch
    scale = min(max_w / width, max_h / height, 1.0)
    return [
        PdfImage(str(path), width=width * scale, height=height * scale),
        Paragraph(caption, styles["Caption"]),
    ]


def build_evidence_pdf(path: Path) -> None:
    styles = pdf_styles()
    story = [
        Paragraph("01 Execution Evidence Report", styles["TitleBlack"]),
        Paragraph("Step-by-step build evidence, screenshots, validation, decisions, and ownership disclosure.", styles["SubtitleMuted"]),
        pdf_meta_table(styles),
        Spacer(1, 10),
        Paragraph("Assessment Requirements Covered", styles["H1Blue"]),
        pdf_bullets(
            styles,
            [
                "Runnable GitHub repository with setup documentation, task folders, screenshots, and report deliverables.",
                "RFQ to CRM automation using n8n as an equivalent orchestration approach.",
                "FastAPI quotation microservice with tests, Docker packaging, and OpenAPI documentation.",
                "Bilingual RAG workflow over English and Arabic sample documents with citations and refusal behavior.",
                "Explicit notes on error handling, maintainability, security hygiene, trade-offs, validation, and AI assistance.",
            ],
        ),
        Paragraph("Work Evidence By Task", styles["H1Blue"]),
        pdf_status_table(styles),
        Paragraph("Task 1 Evidence: RFQ To CRM Automation", styles["H1Blue"]),
        pdf_numbered(
            styles,
            [
                "Created a Docker-based n8n environment with persistent volumes.",
                "Built a webhook intake endpoint for RFQ-style JSON payloads.",
                "Configured OpenAI extraction for urgency, key requirements, and summary.",
                "Merged structured fields with extracted intent and generated CRM metadata.",
                "Persisted CRM records and sales alerts with bilingual reply drafts.",
                "Debugged sandbox, missing-file, and node-reference issues during testing.",
            ],
        ),
        pdf_bullets(
            styles,
            [
                "Trade-off: webhook simulation kept the workflow reproducible without OAuth setup.",
                "Limitation: attachment filenames are captured as metadata; real file storage is a future enhancement.",
                "Validation: multiple full workflow executions produced successful CRM and alert outputs.",
            ],
        ),
    ]
    for name in [
        "repository structure.png",
        "proof of the volumes being correctly attached.png",
        "Webhook node.png",
        "n8n canva.png",
        "open ai node.png",
        "set field.png",
        "set field output.png",
        "javascript code 1 error .png",
        "workflow progress .png",
        "workflow progress 1 .png",
        "workflow.png",
        "workflow2.png",
        "workflow succes.png",
    ]:
        story.extend(pdf_image_flow(SCREENSHOT_DIR / "Task 1" / name, f"Task 1 evidence: {name}", styles))
    story.extend(
        [
            PageBreak(),
            Paragraph("Task 2 Evidence: Quotation Microservice", styles["H1Blue"]),
            pdf_bullets(
                styles,
                [
                    "FastAPI service with health, catalog, quote creation, and quote retrieval endpoints.",
                    "Product catalog and bulk-discount logic with offline-friendly mocked data.",
                    "SQLite persistence, Docker packaging, and pytest coverage.",
                ],
            ),
        ]
    )
    story.extend(pdf_image_flow(SCREENSHOT_DIR / "Task 2" / "Open API.png", "Task 2 evidence: OpenAPI documentation page.", styles))
    story.extend(
        [
            Paragraph("Task 3 Evidence: Bilingual RAG Knowledge Workflow", styles["H1Blue"]),
            pdf_bullets(
                styles,
                [
                    "FastAPI service loads 4 sample documents in English and Arabic.",
                    "Embeddings are cached locally and searched with cosine similarity.",
                    "The service returns cited answers and refuses out-of-scope questions.",
                    "Additional evidence is stored at docs/screenshots/Task 3/testing the RAG on arabic question.pdf.",
                ],
            ),
        ]
    )
    story.extend(
        pdf_image_flow(
            SCREENSHOT_DIR / "Task 3" / "a problem when communicating with model.png",
            "Task 3 evidence: model communication issue captured during testing.",
            styles,
        )
    )
    story.extend(
        [
            Paragraph("Validation Summary", styles["H1Blue"]),
            pdf_bullets(
                styles,
                [
                    "Task 1 was validated through end-to-end n8n executions and persisted mock output files.",
                    "Task 2 includes a pytest suite and OpenAPI evidence.",
                    "Task 3 includes offline tests plus optional API-key integration tests.",
                    "Docker packaging is included for the runtime services.",
                ],
            ),
            Paragraph("Ownership And AI Assistance Disclosure", styles["H1Blue"]),
            Paragraph(
                "AI assistance was used as a technical mentor and documentation partner for architecture discussion, starter logic, debugging support, and final writing polish. The candidate personally completed the hands-on implementation work: n8n node setup, Docker execution, API implementation, test execution, screenshot collection, environment configuration, validation, and final engineering decisions.",
                styles["BodyText"],
            ),
        ]
    )
    pdf_doc(path).build(story)


def build_final_pdf(path: Path) -> None:
    styles = pdf_styles()
    story = [
        Paragraph("02 Final Result Report", styles["TitleBlack"]),
        Paragraph("Final outputs, architecture summary, setup guide, API evidence, limitations, and improvements.", styles["SubtitleMuted"]),
        pdf_meta_table(styles),
        Spacer(1, 10),
        Paragraph("Final Outcome", styles["H1Blue"]),
        Paragraph(
            "The assessment implementation is complete across all three requested tasks. The repository contains runnable source code, Docker packaging, tests where required, exported workflow artifacts, setup documentation, screenshots, and final reports.",
            styles["BodyText"],
        ),
        pdf_status_table(styles),
        Paragraph("Task 1 Final Result", styles["H1Blue"]),
        pdf_bullets(
            styles,
            [
                "Inbound RFQ data enters through an n8n webhook.",
                "OpenAI extracts urgency, key requirements, and summary from the RFQ message.",
                "The workflow creates a mock CRM record, bilingual reply draft, and internal alert.",
            ],
        ),
        Paragraph("Task 2 Final Result", styles["H1Blue"]),
        pdf_bullets(
            styles,
            [
                "FastAPI exposes health, catalog, quote creation, and quote retrieval endpoints.",
                "Bulk discount tiers are applied by total quote quantity.",
                "SQLite stores generated quotes; OpenAPI documentation is available at `/docs`.",
            ],
        ),
        Paragraph("Task 3 Final Result", styles["H1Blue"]),
        pdf_bullets(
            styles,
            [
                "The service indexes 4 English/Arabic source documents.",
                "Responses cite source documents and include latency and estimated cost metadata.",
                "Out-of-scope questions are refused when retrieval confidence is insufficient.",
            ],
        ),
        Paragraph("Architecture Explanation", styles["H1Blue"]),
        pdf_bullets(
            styles,
            [
                "Task 1 uses n8n because the workflow can be exported, versioned, reviewed, and rerun locally.",
                "Task 2 separates schemas, catalog data, pricing logic, database access, and route handlers.",
                "Task 3 keeps the RAG system lightweight because the assessment document set is small.",
            ],
        ),
        Paragraph("Run And Test Commands", styles["H1Blue"]),
    ]
    cmd_data = [
        ["Area", "Command", "Purpose"],
        ["Task 1", "docker-compose up -d", "Start n8n from repository root."],
        ["Task 2", "cd task2_quotation_microservice && docker-compose up --build", "Run quotation API on port 8000."],
        ["Task 2 tests", "cd task2_quotation_microservice && PYTHONPATH=. pytest tests/ -v", "Run quotation tests."],
        ["Task 3", "cd task3_bilingual_rag && docker-compose up --build", "Run RAG API on port 8001."],
        ["Task 3 tests", "cd task3_bilingual_rag && PYTHONPATH=. pytest tests/ -v", "Run RAG tests."],
    ]
    table = Table([[Paragraph(c, styles["BodyText"]) for c in row] for row in cmd_data], colWidths=[1.0 * inch, 3.05 * inch, 2.35 * inch], repeatRows=1)
    table.setStyle(
        TableStyle(
            [
                ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#D0D7DE")),
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#F2F4F7")),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#1F4D78")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 5),
                ("RIGHTPADDING", (0, 0), (-1, -1), 5),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ]
        )
    )
    story.append(table)
    story.extend(
        [
            PageBreak(),
            Paragraph("Repository Guidance", styles["H1Blue"]),
            Paragraph("Use this structure map to navigate the solution bundle and GitHub repository:", styles["BodyText"]),
            Paragraph(repo_tree().replace("\n", "<br/>"), styles["CodeSmall"]),
            Paragraph("Error Handling, Maintainability, And Security", styles["H1Blue"]),
            pdf_bullets(
                styles,
                [
                    "Task 1 persists n8n state and mock CRM data with Docker volumes and bind mounts.",
                    "Task 2 validates input with Pydantic and returns clear HTTP errors.",
                    "Task 3 refuses unsupported questions through retrieval thresholding.",
                    "Secrets are stored in `.env` files or runtime environment variables, not hardcoded.",
                ],
            ),
            Paragraph("Future Enhancements", styles["H1Blue"]),
            pdf_bullets(
                styles,
                [
                    "Task 1: ingest RFQs from real email inboxes, WhatsApp Business, website forms, or CRM webhooks.",
                    "Task 1: archive actual attachment files and add duplicate detection.",
                    "Task 2: add a quotation UI, PDF quote generation, authentication, configurable pricing, taxes, shipping, and ERP/CRM integration.",
                    "Task 3: add a polished bilingual chat UI, document upload, richer citations, streaming responses, access control, and a vector database if the corpus grows.",
                ],
            ),
            Paragraph("AI Assistance Disclosure", styles["H1Blue"]),
            Paragraph(
                "AI assistance was used for brainstorming architecture choices, debugging explanations, starter snippets, and documentation/report polish. The candidate remained responsible for all local execution, implementation, testing, configuration, screenshots, trade-off selection, and final submission packaging.",
                styles["BodyText"],
            ),
        ]
    )
    pdf_doc(path).build(story)


def main() -> None:
    evidence = build_evidence_report()
    final = build_final_report()
    render_docx(evidence)
    render_docx(final)
    print(evidence)
    print(final)


if __name__ == "__main__":
    main()
